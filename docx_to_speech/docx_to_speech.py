"""
AION Plugin: docx_to_speech
============================
Liest eine Word-Datei (.docx) vor und speichert sie als MP3 oder WAV.

Tool:
  - docx_to_speech : .docx → MP3/WAV (kombiniert python-docx + audio_pipeline TTS)

Anforderungen:
  - python-docx   (pip install python-docx)   — Text aus Word-Datei lesen
  - audio_pipeline Plugin aktiv               — TTS-Ausgabe (edge-tts oder sapi5)

Empfohlene Konfiguration in config.json:
  {"tts_engine": "edge", "tts_voice": "de-DE-KatjaNeural"}
"""

import os
from pathlib import Path

from docx import Document


def _extract_docx_text(path: str) -> str:
    """Extrahiert den reinen Text aus einer .docx-Datei (alle Absätze, Tabellen)."""
    doc = Document(path)
    lines = []

    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped:
            lines.append(stripped)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                stripped = cell.text.strip()
                if stripped and stripped not in lines:
                    lines.append(stripped)

    return "\n".join(lines)


def docx_to_speech_func(
    docx_path: str,
    output_path: str = "",
    engine: str = "",
    memory=None,
) -> dict:
    """Liest eine Word-Datei vor und speichert sie als MP3/WAV.

    Importierbar von anderen Plugins:
        from docx_to_speech.docx_to_speech import docx_to_speech_func
        result = docx_to_speech_func("bericht.docx", "bericht.mp3")
    """
    if not docx_path:
        return {"ok": False, "error": "docx_path ist erforderlich"}

    if not os.path.exists(docx_path):
        return {"ok": False, "error": f"Datei nicht gefunden: {docx_path}"}

    if not docx_path.lower().endswith(".docx"):
        return {"ok": False, "error": "Nur .docx-Dateien werden unterstützt"}

    # ── 1. Text extrahieren ───────────────────────────────────────────────────
    try:
        text = _extract_docx_text(docx_path)
    except Exception as e:
        return {"ok": False, "error": f"DOCX lesen fehlgeschlagen: {e}"}

    if not text.strip():
        return {"ok": False, "error": "Dokument enthält keinen lesbaren Text"}

    # ── 2. Output-Pfad ableiten (falls nicht angegeben) ───────────────────────
    if not output_path:
        stem = Path(docx_path).stem
        folder = Path(docx_path).parent
        suffix = ".wav" if engine == "sapi5" else ".mp3"
        output_path = str(folder / (stem + suffix))

    # ── 3. TTS via audio_pipeline ─────────────────────────────────────────────
    try:
        from audio_pipeline.audio_pipeline import audio_tts
    except ImportError:
        return {
            "ok": False,
            "error": (
                "audio_pipeline Plugin nicht gefunden. "
                "Bitte das 'Audio Pipeline' Plugin in AION aktivieren."
            ),
        }

    result = audio_tts(text=text, engine=engine, output_path=output_path)

    if result.get("ok"):
        result["source"] = docx_path
        result["chars"] = len(text)

    return result


def register(api):
    api.register_tool(
        name="docx_to_speech",
        description=(
            "Liest eine Word-Datei (.docx) vor und speichert sie als MP3 oder WAV. "
            "Extrahiert den Text aus dem Dokument und wandelt ihn per TTS in Sprache um. "
            "Benötigt das Audio Pipeline Plugin. "
            "Gibt {ok, path, engine, chars, source} zurück."
        ),
        func=docx_to_speech_func,
        input_schema={
            "type": "object",
            "properties": {
                "docx_path": {
                    "type": "string",
                    "description": "Vollständiger Pfad zur .docx-Datei (z. B. C:/Dokumente/Bericht.docx)",
                },
                "output_path": {
                    "type": "string",
                    "description": (
                        "Ausgabepfad für die Audiodatei (z. B. C:/Dokumente/Bericht.mp3). "
                        "Leer = wird automatisch neben der Quelldatei generiert."
                    ),
                },
                "engine": {
                    "type": "string",
                    "description": (
                        "TTS-Engine: 'edge' (Microsoft Neural, online, beste Qualität) "
                        "oder 'sapi5' (offline, Fallback). "
                        "Leer = Einstellung aus config.json."
                    ),
                },
            },
            "required": ["docx_path"],
        },
    )
